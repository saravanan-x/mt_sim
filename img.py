from PIL import Image
import numpy as np
from docx import Document

# Load image
image_path = "fish.jpeg"
img = Image.open(image_path).convert("L")

# Convert to numpy array
img_array = np.array(img)

# Convert to binary
binary = (img_array > 128).astype(int)

# Convert binary to hex
hex_data = []

for row in binary:
    bits = ''.join(map(str, row))

    for i in range(0, len(bits), 8):
        byte = bits[i:i+8]

        if len(byte) == 8:
            hex_value = format(int(byte, 2), '02X')
            hex_data.append("0x" + hex_value)

# Create Word document
doc = Document()
doc.add_heading("Image Hex Data", level=1)

# Convert list to series
hex_series = ", ".join(hex_data)

doc.add_paragraph(hex_series)

# Save Word file
doc.save("image_hex_series.docx")

print("Word file saved with hex values in series.")